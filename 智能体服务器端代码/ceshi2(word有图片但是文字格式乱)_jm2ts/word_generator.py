"""
Word文档生成器
支持图片插入的增强版markdown转word功能
"""

import os
import re
import uuid
import logging
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import markdown
from bs4 import BeautifulSoup
from markdown_processor import MarkdownProcessor

logger = logging.getLogger(__name__)

class WordGenerator:
    """Word文档生成器，支持图片插入 - 强制字体设置版本"""

    def __init__(self, output_dir: str = "/var/www/files"):
        """
        初始化Word生成器

        Args:
            output_dir: 输出目录
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.markdown_processor = MarkdownProcessor()

        # 强制字体配置 - 解决字体格式混乱问题
        self.fonts = {
            'chinese': '微软雅黑',      # 中文字体
            'english': 'Calibri',       # 英文字体
            'code': 'Consolas',         # 代码字体
            'fallback': '宋体'          # 备用字体
        }

        # 字号配置
        self.font_sizes = {
            'title': 22,        # 主标题
            'heading1': 18,     # 一级标题
            'heading2': 16,     # 二级标题
            'heading3': 14,     # 三级标题
            'normal': 12,       # 正文
            'small': 10,        # 小字
            'code': 10          # 代码
        }
    
    def convert_markdown_to_docx(self, markdown_content: str, 
                                chart_configs: Dict[str, Dict] = None,
                                filename: str = None) -> str:
        """
        将markdown内容转换为Word文档，支持图片插入
        
        Args:
            markdown_content: markdown内容
            chart_configs: 图表配置字典
            filename: 输出文件名，如果为None则自动生成
            
        Returns:
            生成的Word文档路径
        """
        if filename is None:
            filename = f"doc_{uuid.uuid4().hex[:8]}.docx"
        
        output_path = os.path.join(self.output_dir, filename)
        
        try:
            # 处理markdown中的图表标记
            processed_md, image_paths = self.markdown_processor.process_markdown_with_charts(
                markdown_content, chart_configs
            )
            
            # 创建Word文档
            doc = Document()
            
            # 强制设置文档样式和字体
            self._setup_document_styles(doc)
            self._force_document_fonts(doc)
            
            # 解析markdown为HTML
            html_content = markdown.markdown(processed_md, extensions=['tables', 'fenced_code'])
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 转换HTML元素为Word元素
            self._convert_html_to_docx(doc, soup, image_paths)
            
            # 保存文档
            doc.save(output_path)
            
            # 清理临时图片
            self.markdown_processor.clean_temp_images()
            
            logger.info(f"成功生成Word文档: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"生成Word文档失败: {str(e)}")
            raise
    
    def _setup_document_styles(self, doc: Document):
        """强制设置文档样式和字体"""
        # 设置Normal样式 - 强制字体
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.fonts['chinese']
        normal_font.size = Pt(self.font_sizes['normal'])
        normal_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # 设置段落格式
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.space_after = Pt(6)
        normal_paragraph.line_spacing = 1.15
        normal_paragraph.first_line_indent = Pt(0)

        # 强制设置英文字体
        self._set_font_for_complex_script(normal_style, self.fonts['english'])

        # 设置标题样式 - 精确控制字体格式
        heading_configs = [
            ('Heading 1', self.font_sizes['heading1'], False),  # 一级标题不加粗
            ('Heading 2', self.font_sizes['heading2'], False),  # 二级标题不加粗
            ('Heading 3', self.font_sizes['heading3'], False),  # 三级标题不加粗
            ('Heading 4', 13, False),
            ('Heading 5', 12, False),
            ('Heading 6', 11, False),
        ]

        for style_name, font_size, is_bold in heading_configs:
            try:
                heading_style = doc.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = self.fonts['chinese']
                heading_font.size = Pt(font_size)
                heading_font.bold = is_bold  # 明确设置是否加粗
                heading_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                # 设置标题段落格式
                heading_paragraph = heading_style.paragraph_format
                heading_paragraph.space_before = Pt(12)
                heading_paragraph.space_after = Pt(6)
                heading_paragraph.keep_with_next = True

                # 强制设置复杂脚本字体
                self._set_font_for_complex_script(heading_style, self.fonts['english'])

                # 额外强制设置标题字体属性
                self._force_heading_font_properties(heading_style, font_size, is_bold)

            except KeyError:
                continue

    def _force_document_fonts(self, doc: Document):
        """强制设置文档字体主题"""
        try:
            # 这个方法确保文档级别的字体设置
            # 主要是为了兼容性，确保字体在所有环境下都能正确显示
            pass
        except Exception as e:
            logger.warning(f"设置文档字体主题失败: {e}")

    def _set_font_for_complex_script(self, style, font_name):
        """为样式设置复杂脚本字体（处理中英文混合）"""
        try:
            # 获取样式的rPr元素
            style_element = style.element
            rPr = style_element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style_element.insert(0, rPr)

            # 设置字体
            fonts_element = rPr.find(qn('w:rFonts'))
            if fonts_element is None:
                fonts_element = OxmlElement('w:rFonts')
                rPr.append(fonts_element)

            # 设置各种字体属性
            fonts_element.set(qn('w:ascii'), self.fonts['english'])      # ASCII字符
            fonts_element.set(qn('w:hAnsi'), self.fonts['english'])      # 高位ANSI字符
            fonts_element.set(qn('w:eastAsia'), self.fonts['chinese'])   # 东亚字符
            fonts_element.set(qn('w:cs'), self.fonts['chinese'])         # 复杂脚本字符

        except Exception as e:
            logger.warning(f"设置复杂脚本字体失败: {e}")

    def _force_heading_font_properties(self, style, font_size, is_bold):
        """强制设置标题字体属性"""
        try:
            # 获取样式的rPr元素
            style_element = style.element
            rPr = style_element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style_element.insert(0, rPr)

            # 强制设置字体大小
            sz_element = rPr.find(qn('w:sz'))
            if sz_element is None:
                sz_element = OxmlElement('w:sz')
                rPr.append(sz_element)
            sz_element.set(qn('w:val'), str(font_size * 2))  # Word中字号是磅值的2倍

            # 强制设置字体大小（复杂脚本）
            szCs_element = rPr.find(qn('w:szCs'))
            if szCs_element is None:
                szCs_element = OxmlElement('w:szCs')
                rPr.append(szCs_element)
            szCs_element.set(qn('w:val'), str(font_size * 2))

            # 强制设置是否加粗
            b_element = rPr.find(qn('w:b'))
            bCs_element = rPr.find(qn('w:bCs'))

            if is_bold:
                if b_element is None:
                    b_element = OxmlElement('w:b')
                    rPr.append(b_element)
                b_element.set(qn('w:val'), '1')

                if bCs_element is None:
                    bCs_element = OxmlElement('w:bCs')
                    rPr.append(bCs_element)
                bCs_element.set(qn('w:val'), '1')
            else:
                # 明确设置不加粗
                if b_element is None:
                    b_element = OxmlElement('w:b')
                    rPr.append(b_element)
                b_element.set(qn('w:val'), '0')

                if bCs_element is None:
                    bCs_element = OxmlElement('w:bCs')
                    rPr.append(bCs_element)
                bCs_element.set(qn('w:val'), '0')

            # 强制设置字体颜色为黑色
            color_element = rPr.find(qn('w:color'))
            if color_element is None:
                color_element = OxmlElement('w:color')
                rPr.append(color_element)
            color_element.set(qn('w:val'), '000000')

        except Exception as e:
            logger.warning(f"强制设置标题字体属性失败: {e}")

    def _force_heading_paragraph_font(self, paragraph, level):
        """强制设置标题段落的字体格式"""
        try:
            # 确定字体大小
            font_size_map = {
                1: self.font_sizes['heading1'],
                2: self.font_sizes['heading2'],
                3: self.font_sizes['heading3'],
                4: 13,
                5: 12,
                6: 11
            }
            font_size = font_size_map.get(level, 12)

            # 强制设置段落中所有run的字体
            for run in paragraph.runs:
                # 设置基本字体属性
                run.font.name = self.fonts['chinese']
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.bold = False  # 标题不加粗
                run.italic = False

                # XML级别强制设置
                self._force_run_complex_font(run, is_bold=False, is_italic=False)

            # 如果段落没有run，创建一个
            if not paragraph.runs:
                run = paragraph.add_run("")
                run.font.name = self.fonts['chinese']
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.bold = False
                run.italic = False
                self._force_run_complex_font(run, is_bold=False, is_italic=False)

        except Exception as e:
            logger.warning(f"强制设置标题段落字体失败: {e}")
    
    def _convert_html_to_docx(self, doc: Document, soup: BeautifulSoup, 
                            image_paths: Dict[str, str]):
        """将HTML内容转换为Word文档元素"""
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'img', 'ul', 'ol']):
            if element.name.startswith('h'):
                # 处理标题 - 强制字体设置
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text().strip(), level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 强制设置标题字体格式
                self._force_heading_paragraph_font(heading, level)
                
            elif element.name == 'p':
                # 处理段落
                paragraph = doc.add_paragraph()
                self._process_paragraph_content(paragraph, element)
                
            elif element.name == 'table':
                # 处理表格
                self._add_table_to_doc(doc, element)
                
            elif element.name == 'img':
                # 处理图片
                self._add_image_to_doc(doc, element, image_paths)
                
            elif element.name in ['ul', 'ol']:
                # 处理列表
                self._add_list_to_doc(doc, element)
    
    def _process_paragraph_content(self, paragraph, p_element):
        """处理段落内容，强制字体设置"""
        # 设置段落基本格式
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 如果段落为空，添加一个空的run
        if not p_element.contents:
            run = paragraph.add_run("")
            self._force_run_font(run)
            return

        # 处理段落内容
        for content in p_element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong' or content.name == 'b':
                    # 粗体
                    run = paragraph.add_run(content.get_text())
                    self._force_run_font(run, is_bold=True)
                elif content.name == 'em' or content.name == 'i':
                    # 斜体
                    run = paragraph.add_run(content.get_text())
                    self._force_run_font(run, is_italic=True)
                elif content.name == 'code':
                    # 代码
                    run = paragraph.add_run(content.get_text())
                    self._force_run_font(run, is_code=True)
                else:
                    # 其他标签
                    run = paragraph.add_run(content.get_text())
                    self._force_run_font(run)
            else:
                # 纯文本
                text = str(content).strip()
                if text:
                    run = paragraph.add_run(text)
                    self._force_run_font(run)

    def _force_run_font(self, run, is_bold=False, is_italic=False, is_code=False, is_heading=False):
        """强制设置文本运行的字体格式 - 精确控制加粗"""
        if is_code:
            run.font.name = self.fonts['code']
            run.font.size = Pt(self.font_sizes['code'])
            run.font.color.rgb = RGBColor(0x2F, 0x4F, 0x4F)
        elif is_heading:
            run.font.name = self.fonts['chinese']
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        else:
            run.font.name = self.fonts['chinese']
            run.font.size = Pt(self.font_sizes['normal'])
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # 精确设置格式 - 明确控制加粗状态
        run.bold = is_bold  # 明确设置加粗状态，而不是累加
        run.italic = is_italic  # 明确设置斜体状态

        # 强制设置复杂脚本字体
        self._force_run_complex_font(run, is_bold, is_italic)

    def _force_run_complex_font(self, run, is_bold=False, is_italic=False):
        """强制设置run的复杂脚本字体和格式"""
        try:
            # 获取run的rPr元素
            run_element = run.element
            rPr = run_element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run_element.insert(0, rPr)

            # 设置字体
            fonts_element = rPr.find(qn('w:rFonts'))
            if fonts_element is None:
                fonts_element = OxmlElement('w:rFonts')
                rPr.append(fonts_element)

            # 强制设置所有字体属性
            fonts_element.set(qn('w:ascii'), self.fonts['english'])
            fonts_element.set(qn('w:hAnsi'), self.fonts['english'])
            fonts_element.set(qn('w:eastAsia'), self.fonts['chinese'])
            fonts_element.set(qn('w:cs'), self.fonts['chinese'])

            # 强制设置加粗属性
            b_element = rPr.find(qn('w:b'))
            bCs_element = rPr.find(qn('w:bCs'))

            if is_bold:
                if b_element is None:
                    b_element = OxmlElement('w:b')
                    rPr.append(b_element)
                b_element.set(qn('w:val'), '1')

                if bCs_element is None:
                    bCs_element = OxmlElement('w:bCs')
                    rPr.append(bCs_element)
                bCs_element.set(qn('w:val'), '1')
            else:
                # 明确设置不加粗
                if b_element is None:
                    b_element = OxmlElement('w:b')
                    rPr.append(b_element)
                b_element.set(qn('w:val'), '0')

                if bCs_element is None:
                    bCs_element = OxmlElement('w:bCs')
                    rPr.append(bCs_element)
                bCs_element.set(qn('w:val'), '0')

            # 强制设置斜体属性
            i_element = rPr.find(qn('w:i'))
            iCs_element = rPr.find(qn('w:iCs'))

            if is_italic:
                if i_element is None:
                    i_element = OxmlElement('w:i')
                    rPr.append(i_element)
                i_element.set(qn('w:val'), '1')

                if iCs_element is None:
                    iCs_element = OxmlElement('w:iCs')
                    rPr.append(iCs_element)
                iCs_element.set(qn('w:val'), '1')
            else:
                # 明确设置不斜体
                if i_element is None:
                    i_element = OxmlElement('w:i')
                    rPr.append(i_element)
                i_element.set(qn('w:val'), '0')

                if iCs_element is None:
                    iCs_element = OxmlElement('w:iCs')
                    rPr.append(iCs_element)
                iCs_element.set(qn('w:val'), '0')

        except Exception as e:
            logger.warning(f"强制设置run字体失败: {e}")

    # 保留原有方法名以兼容现有代码
    def _set_run_font(self, run, is_bold=False, is_italic=False, is_code=False, is_normal=False):
        """兼容性方法 - 调用新的强制字体设置方法"""
        self._force_run_font(run, is_bold=is_bold, is_italic=is_italic, is_code=is_code)
    
    def _add_table_to_doc(self, doc: Document, table_element):
        """添加表格到文档 - 强制字体设置"""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # 获取列数
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充表格内容
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < len(table.rows[i].cells):
                    table_cell = table.rows[i].cells[j]
                    cell_text = cell.get_text().strip()

                    # 清空单元格默认内容
                    table_cell.text = ""

                    # 添加段落和文本
                    paragraph = table_cell.paragraphs[0]
                    run = paragraph.add_run(cell_text)

                    # 强制设置字体
                    if cell.name == 'th':
                        self._force_run_font(run, is_bold=True)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        self._force_run_font(run)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # 设置段落格式
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
    
    def _add_image_to_doc(self, doc: Document, img_element, image_paths: Dict[str, str]):
        """添加图片到文档 - 强制字体设置"""
        src = img_element.get('src', '')
        alt_text = img_element.get('alt', '图片')

        if src in image_paths:
            image_path = image_paths[src]

            if os.path.exists(image_path):
                try:
                    # 添加图片标题
                    caption_paragraph = doc.add_paragraph()
                    caption_run = caption_paragraph.add_run(alt_text)
                    self._force_run_font(caption_run, is_bold=True)
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 添加图片
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                    # 设置图片大小（最大宽度6英寸）
                    run.add_picture(image_path, width=Inches(6))

                    logger.info(f"成功插入图片: {image_path}")

                except Exception as e:
                    logger.error(f"插入图片失败 {image_path}: {str(e)}")
                    # 如果图片插入失败，添加文本说明
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(f"[图片: {alt_text}]")
                    self._force_run_font(run)
            else:
                logger.warning(f"图片文件不存在: {image_path}")
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"[图片: {alt_text} - 文件不存在]")
                self._force_run_font(run)
        else:
            # 处理file://路径或其他类型的图片引用
            if src.startswith('file://'):
                image_path = src[7:]  # 移除file://前缀

                if os.path.exists(image_path):
                    try:
                        # 添加图片标题
                        caption_paragraph = doc.add_paragraph()
                        caption_run = caption_paragraph.add_run(alt_text)
                        self._force_run_font(caption_run, is_bold=True)
                        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # 添加图片
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                        # 设置图片大小（最大宽度6英寸）
                        run.add_picture(image_path, width=Inches(6))

                        logger.info(f"成功插入图片: {image_path}")

                    except Exception as e:
                        logger.error(f"插入图片失败 {image_path}: {str(e)}")
                        # 如果图片插入失败，添加文本说明
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(f"[图片: {alt_text}]")
                        self._force_run_font(run)
                else:
                    logger.warning(f"图片文件不存在: {image_path}")
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(f"[图片: {alt_text} - 文件不存在]")
                    self._force_run_font(run)
            else:
                # 其他类型的图片引用
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"[图片: {alt_text}]")
                self._force_run_font(run)

    def _add_list_to_doc(self, doc: Document, list_element):
        """添加列表到文档 - 强制字体设置"""
        list_items = list_element.find_all('li')

        for item in list_items:
            paragraph = doc.add_paragraph()

            # 设置列表样式
            if list_element.name == 'ul':
                # 无序列表
                paragraph.style = 'List Bullet'
            else:
                # 有序列表
                paragraph.style = 'List Number'

            # 添加文本并强制设置字体
            run = paragraph.add_run(item.get_text().strip())
            self._force_run_font(run)

            # 设置段落格式
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15

def convert_markdown_to_word_with_charts(markdown_content: str,
                                       chart_configs: Dict[str, Dict] = None,
                                       output_dir: str = "/var/www/files",
                                       filename: str = None) -> str:
    """
    便捷函数：将包含图表的markdown转换为Word文档
    
    Args:
        markdown_content: markdown内容
        chart_configs: 图表配置字典
        output_dir: 输出目录
        filename: 文件名
        
    Returns:
        生成的Word文档路径
    """
    generator = WordGenerator(output_dir)
    return generator.convert_markdown_to_docx(markdown_content, chart_configs, filename)

# 测试代码
if __name__ == "__main__":
    test_markdown = """
# 垃圾焚烧企业数据分析报告

## 1. 数据概览

| 指标 | 数值 | 单位 |
|------|------|------|
| 处理量 | 1200 | 吨 |
| 运行时间 | 720 | 小时 |
| 平均温度 | 865 | ℃ |

## 2. 温度趋势分析

![炉膛温度变化趋势](chart_temp)

温度控制在**正常范围**内，符合环保要求。

## 3. 结论

- 设备运行稳定
- 环保指标达标
- 建议继续保持现有操作水平
"""
    
    chart_configs = {
        "chart_temp": {
            "title": { "text": "炉膛温度变化趋势" },
            "xAxis": { 
                "type": "category", 
                "data": ["1月", "2月", "3月", "4月", "5月", "6月"] 
            },
            "yAxis": { "type": "value" },
            "series": [{
                "type": "line",
                "name": "温度",
                "data": [850, 860, 855, 870, 865, 875]
            }]
        }
    }
    
    # 生成Word文档
    output_path = convert_markdown_to_word_with_charts(
        test_markdown, 
        chart_configs,
        output_dir="/tmp",
        filename="test_report.docx"
    )
    
    print(f"生成的Word文档: {output_path}")
